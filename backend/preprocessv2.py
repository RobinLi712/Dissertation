import pandas as pd
import docx
import pytesseract
from PIL import Image
import io
import os

pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_image(image):
    return pytesseract.image_to_string(image)

def extract_images_from_docx(file_path):
    doc = docx.Document(file_path)
    images_texts = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image = doc.part.rels[rel.rId].target_part.blob
            image = Image.open(io.BytesIO(image))
            text = extract_text_from_image(image)
            images_texts.append(text)
    return '\n'.join(images_texts)

def read_full_docx(file_path):
    text_content = read_docx(file_path)
    image_content = extract_images_from_docx(file_path)
    full_text = text_content + '\n' + image_content
    print(f"Full text from {file_path}: {full_text[:500]}")  # 打印前500个字符作为示例
    return full_text

def clean_text(text):
    text = text.strip().replace('\n', ' ')
    print(f"Cleaned text: {text[:200]}")  # 打印前200个字符作为示例
    return text

def split_text(text, max_length=512):
    # 按max_length字符切割文本
    return [text[i:i+max_length] for i in range(0, len(text), max_length)]

def create_dataset(excel_path, dataset_dir, output_csv):
    data = pd.read_excel(excel_path)
    keywords = data['Content Summary (key words)'].values
    links = data['Link to Response'].values

    dataset = []
    for keyword, link in zip(keywords, links):
        try:
            file_path = os.path.join(dataset_dir, link)
            content = read_full_docx(file_path)
            cleaned_content = clean_text(content)
            split_contents = split_text(cleaned_content, max_length=512)
            for split_content in split_contents:
                dataset.append([keyword, split_content])
        except Exception as e:
            print(f"Error reading {link}: {e}")

    dataset_df = pd.DataFrame(dataset, columns=['Keywords', 'Content'])
    dataset_df.to_csv(output_csv, index=False)

if __name__ == "__main__":
    excel_path = '../dataset/bid.xlsx'
    dataset_dir = '../dataset'
    output_csv = 'distilbert_dataset.csv'
    create_dataset(excel_path, dataset_dir, output_csv)
