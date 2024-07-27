import pandas as pd
import docx
import pytesseract
from PIL import Image
import io
import os

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_image(image):
    return pytesseract.image_to_string(image)

def extract_images_from_docx(file_path):
    doc = docx.Document(file_path)
    images_texts = []
    for rel in doc.part.rels.values():
        if "image" in doc.part.rels[rel.rId].target_ref:
            image = doc.part.rels[rel.rId].target_part.blob
            image = Image.open(io.BytesIO(image))
            text = extract_text_from_image(image)
            images_texts.append(text)
        return '\n'.join(images_texts)
    
def read_full_docx(file_path):
    text_content = read_docx(file_path)
    image_content = extract_images_from_docx(file_path)
    return text_content + '\n' + image_content

def create_dataset(excel_path, dataset_dir, output_csv):
    # Read the excel file
    data = pd.read_excel(excel_path)

    # Keyword extraction
    keywords = data['Content Summary (key words)'].values
    links = data['Link to Response'].values

    # Build the dataset
    data = []
    for keyword, link in zip(keywords, links):
        try:
            file_path = os.path.join(dataset_dir, link)
            content = read_full_docx(file_path)
            data.append([keyword, content])
        except Exception as e:
            print(f"Error reading {link}: {e}")

    # transform the data into a DataFrame
    dataset = pd.DataFrame(data, columns=['Keywords', 'Content'])

    # Save the dataset
    dataset.to_csv(output_csv, index=False)

if __name__ == "__main__":
    excel_path = '../dataset/bid.xlsx'
    dataset_dir = '../dataset'
    output_csv = 'dataset.csv'
    create_dataset(excel_path, dataset_dir, output_csv)